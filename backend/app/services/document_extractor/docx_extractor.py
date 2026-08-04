from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore

from app.services.document_extractor.exceptions import CorruptedDocumentError


class DocxExtractor:
    """Extract text from DOCX while preserving reading order."""

    def extract(self, file_path: str | Path) -> str:
        try:
            document = Document(str(file_path))
        except Exception as exc:  # pragma: no cover - dependency/runtime path
            raise CorruptedDocumentError("Unable to read the DOCX file") from exc

        paragraphs: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        for table in document.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                paragraphs.extend(rows)

        return "\n".join(paragraphs)
